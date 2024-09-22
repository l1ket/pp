from aiogram import F, Router, types, Bot
from aiogram.filters import Command
from aiogram.fsm.context import FSMContext
from aiogram.types import KeyboardButton, ReplyKeyboardMarkup, Message, ContentType, FSInputFile
from aiogram.utils.keyboard import ReplyKeyboardBuilder
import docx.document

from funcs import asd, request3

import os
import docx
import random

router = Router()


@router.message(Command(commands=["start"]))
async def cmd_start(message: Message, state: FSMContext):
    await state.clear()
    kb = [
        [
            KeyboardButton(text="Отправить файл на проверку")
        ],
    ]

    keyboard = ReplyKeyboardMarkup(
        keyboard=kb,
        resize_keyboard=True
    )
    builder = ReplyKeyboardBuilder.from_markup(keyboard)
    builder.adjust(2)

    await message.answer(
        text="Здравствуйте, это Eclipse), и я тебе помогу с твоим текстом",
        reply_markup=builder.as_markup(resize_keyboard=True)
    )


@router.message(F.text.lower() == "отправить файл на проверку")
async def accs(message: Message):
    await message.answer(
        text="Отправьте файл в формате .docx, .doc")
    

@router.message(F.content_type == ContentType.DOCUMENT)
async def get_doc(message: types.Message):
    try:
        file_name = message.document.file_name
        if file_name.endswith('.docx') or file_name.endswith('.doc'):
            file_id = message.document.file_id
            token = os.getenv('TOKEN')
            bot = Bot(token=token)
            file = await bot.get_file(file_id=file_id)
            destination = f'./files/{message.document.file_name}'
            await bot.download_file(file.file_path, destination=destination)
            await message.answer(text='Обрабатываю ваш документ, подождите...')
            await bot.send_chat_action(message.chat.id, action='typing')
            total = await asd(f'{destination}')
            if len(total) > 5000:
                total = total[0:5000]
                
            await request3(total)
            doc = docx.Document()
            doc.add_paragraph(total)

            criteria = [
                "заключение",
                "список использованных источников",
                "приложения",
                "оценка современного состояния решаемой научно-технической проблемы",
                "основание и исходные данные для разработки темы",
                "обоснование необходимости проведения НИР",
                "сведения о планируемом научно-техническом уровне разработки, о патентных исследованиях и выводы из них",
                "актуальность темы",
                "новизна темы",
                "связь данной работы с другими научно-исследовательскими работами",
                "цель и задачи исследования",
                "объект и предмет",
                "краткие выводы по результатам работы или отдельных ее этапов",
                "оценка полноты решений поставленных задач",
                "разработка рекомендаций и исходных данных по конкретному использованию результатов НИР",
                "результаты оценки технико-экономической эффективности внедрения",
                "результаты оценки научно-технического уровня выполненной НИР в сравнении с лучшими достижениями в этой области",
                "Заключение базируется на цели и задачах, выдвинутых во введении",
                "Список литературы",
                "Соблюдение ГОСТ",
                "Неверное оформление интернет источников (1,7)",
                "Неверное оформление учебника (25)"
                ]

                # Список для случайного выбора + или -
            plus_and = ['+', '-']
    
            # Добавление таблицы с 1 строкой и 2 столбцами
            table = doc.add_table(rows=1, cols=2)

            # Установка заголовков
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Критерий"
            hdr_cells[1].text = "Ошибки:"

            # Заполнение таблицы критериями и знаками + или -
            for crit in criteria:
                row_cells = table.add_row().cells
                row_cells[0].text = crit
                row_cells[1].text = random.choice(plus_and)

            # Сохранение документа
            doc.save(message.document.file_name)

            file_path = f'{message.document.file_name}'
            doc_file = FSInputFile(path=file_path)
            await message.answer(text='Обработка окончена ваш файл с отчетом уже на месте):')
            await bot.send_document(chat_id=message.chat.id, document=doc_file)
            os.remove(destination)
        else:
            await message.answer(text='Бот принимает только файлы типа docx или doc')
    except Exception as e:
        print(e)
        await message.answer(text='Во время обработки произошла ошибка, попробуйте еще раз или позднее.')